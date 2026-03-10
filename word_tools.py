"""
Word操作ツール群 - python-docxを使用してWordファイルを操作する
"""
import os
import re
import json
from typing import Any, Optional
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# 現在開いているドキュメントのキャッシュ
_doc_cache: dict[str, Document] = {}

# 操作を許可するベースディレクトリ（起動時のカレントディレクトリ）
_BASE_DIR = os.path.realpath(os.getcwd())

_HEX_COLOR_RE = re.compile(r'^[0-9A-Fa-f]{6}$')


def _reraise_if_fatal(e: Exception) -> None:
    """KeyboardInterrupt / SystemExit は握りつぶさず再 raise する"""
    if isinstance(e, (KeyboardInterrupt, SystemExit)):
        raise


def _safe_path(file_path: str) -> str:
    """パストラバーサルを防ぐ。BASE_DIR 外のパスは拒否する。"""
    abs_path = os.path.realpath(os.path.join(_BASE_DIR, file_path))
    if not abs_path.startswith(_BASE_DIR + os.sep) and abs_path != _BASE_DIR:
        raise PermissionError(
            f"アクセス拒否: 作業ディレクトリ外のパスは操作できません: {abs_path}"
        )
    return abs_path


def _validate_hex_color(value: str, name: str) -> str:
    normalized = value.lstrip("#").upper()
    if not _HEX_COLOR_RE.match(normalized):
        raise ValueError(f"{name} は6桁の16進数カラーコードで指定してください（例: FF0000）")
    return normalized


def _get_doc(file_path: str, create_if_missing: bool = False) -> Document:
    abs_path = _safe_path(file_path)
    if abs_path in _doc_cache:
        return _doc_cache[abs_path]
    if os.path.exists(abs_path):
        doc = Document(abs_path)
    elif create_if_missing:
        doc = Document()
    else:
        raise FileNotFoundError(f"ファイルが見つかりません: {abs_path}")
    _doc_cache[abs_path] = doc
    return doc


def _save(file_path: str):
    abs_path = _safe_path(file_path)
    doc = _doc_cache.get(abs_path)
    if doc:
        doc.save(abs_path)


def _para_summary(para, index: int) -> dict:
    """段落の要約情報を返す"""
    return {
        "index": index,
        "style": para.style.name,
        "text": para.text,
        "bold": any(r.bold for r in para.runs if r.bold),
        "italic": any(r.italic for r in para.runs if r.italic),
    }


# ── ツール関数 ──────────────────────────────────────────────────────────────


def open_word(file_path: str, create_if_missing: bool = False) -> dict:
    """Wordファイルを開く"""
    try:
        doc = _get_doc(file_path, create_if_missing=create_if_missing)
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        return {
            "success": True,
            "file_path": os.path.abspath(file_path),
            "paragraph_count": para_count,
            "table_count": table_count,
            "message": f"ドキュメントを開きました。段落数: {para_count}, テーブル数: {table_count}"
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def read_document(file_path: str, include_tables: bool = True) -> dict:
    """ドキュメント全体の内容を読み取る（段落インデックス付き）"""
    try:
        doc = _get_doc(file_path)
        paragraphs = [_para_summary(p, i) for i, p in enumerate(doc.paragraphs)]

        # 全文テキスト
        full_text = "\n".join(p["text"] for p in paragraphs)

        tables = []
        if include_tables:
            for t_idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                tables.append({"table_index": t_idx, "rows": rows})

        return {
            "success": True,
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
            "full_text": full_text,
            "tables": tables,
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def read_paragraph(file_path: str, index: int) -> dict:
    """特定インデックスの段落を読み取る"""
    try:
        doc = _get_doc(file_path)
        if index < 0 or index >= len(doc.paragraphs):
            return {"success": False, "error": f"インデックス {index} は範囲外です（0〜{len(doc.paragraphs)-1}）"}
        para = doc.paragraphs[index]
        runs = [{"text": r.text, "bold": r.bold, "italic": r.italic,
                  "font_size": r.font.size.pt if r.font.size else None,
                  "font_color": str(r.font.color.rgb) if r.font.color and r.font.color.type and r.font.color.rgb else None}
                for r in para.runs]
        return {
            "success": True,
            "index": index,
            "style": para.style.name,
            "text": para.text,
            "runs": runs,
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def append_paragraph(file_path: str, text: str,
                     style: Optional[str] = None,
                     bold: bool = False,
                     italic: bool = False,
                     font_size: Optional[int] = None) -> dict:
    """ドキュメントの末尾に段落を追加する"""
    try:
        doc = _get_doc(file_path, create_if_missing=True)
        para = doc.add_paragraph(style=style)
        run = para.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if font_size:
            run.font.size = Pt(font_size)
        _save(file_path)
        new_index = len(doc.paragraphs) - 1
        return {
            "success": True,
            "message": f"段落を末尾（インデックス {new_index}）に追加しました",
            "index": new_index,
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def append_rich_paragraph(file_path: str,
                           runs: list[dict],
                           style: Optional[str] = None,
                           font_size: Optional[int] = None) -> dict:
    """
    書式付きランのリストから段落を末尾に追加する。
    runs の各要素: {"text": str, "bold": bool, "italic": bool, "color": str}
    color は FF0000 のような6桁HEXまたは #FF0000 形式。
    parse_inline_formatting / parse_markdown の runs フィールドをそのまま渡せる。
    """
    try:
        doc = _get_doc(file_path, create_if_missing=True)
        para = doc.add_paragraph(style=style)
        for r in runs:
            run = para.add_run(r.get("text", ""))
            if r.get("bold"):
                run.bold = True
            if r.get("italic"):
                run.italic = True
            run_size = r.get("font_size") or font_size
            if run_size:
                run.font.size = Pt(run_size)
            if r.get("color"):
                hex_color = _validate_hex_color(r["color"], "color")
                run.font.color.rgb = RGBColor(
                    int(hex_color[0:2], 16),
                    int(hex_color[2:4], 16),
                    int(hex_color[4:6], 16),
                )
        _save(file_path)
        new_index = len(doc.paragraphs) - 1
        return {
            "success": True,
            "message": f"書式付き段落をインデックス {new_index} に追加しました",
            "index": new_index,
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def insert_paragraph(file_path: str, index: int, text: str,
                     position: str = "before",
                     style: Optional[str] = None,
                     bold: bool = False,
                     italic: bool = False,
                     font_size: Optional[int] = None) -> dict:
    """指定インデックスの段落の前後にテキストを挿入する
    position: 'before' または 'after'
    """
    try:
        doc = _get_doc(file_path, create_if_missing=True)
        paras = doc.paragraphs

        if index < 0 or index >= len(paras):
            return {"success": False, "error": f"インデックス {index} は範囲外です（0〜{len(paras)-1}）"}

        ref_para = paras[index]

        # 新しい段落要素を作成
        new_para_elem = OxmlElement("w:p")
        if position == "after":
            ref_para._p.addnext(new_para_elem)
        else:
            ref_para._p.addprevious(new_para_elem)

        # 挿入された段落を探す（XMLから直接特定）
        inserted_index = index if position == "before" else index + 1
        target_para = doc.paragraphs[inserted_index]

        # スタイル設定
        if style:
            target_para.style = doc.styles[style]

        # テキストとフォーマットを設定
        run = target_para.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if font_size:
            run.font.size = Pt(font_size)

        _save(file_path)
        return {
            "success": True,
            "message": f"インデックス {index} の{('前' if position == 'before' else '後')}に段落を挿入しました（新インデックス: {inserted_index}）",
            "inserted_index": inserted_index,
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def replace_text(file_path: str, old_text: str, new_text: str,
                 all_occurrences: bool = True) -> dict:
    """ドキュメント内のテキストを検索・置換する"""
    try:
        doc = _get_doc(file_path)
        count = 0

        for para in doc.paragraphs:
            if old_text not in para.text:
                continue
            # Run単位で置換（書式を保持するためrun内で処理）
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(
                        old_text, new_text, -1 if all_occurrences else 1
                    )
                    count += 1
                    if not all_occurrences:
                        break
            if not all_occurrences and count > 0:
                break

        # テーブル内も処理
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1

        _save(file_path)
        return {
            "success": True,
            "replaced_count": count,
            "message": f"「{old_text}」→「{new_text}」に {count} 箇所置換しました"
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def delete_paragraph(file_path: str, index: int) -> dict:
    """指定インデックスの段落を削除する"""
    try:
        doc = _get_doc(file_path)
        if index < 0 or index >= len(doc.paragraphs):
            return {"success": False, "error": f"インデックス {index} は範囲外です（0〜{len(doc.paragraphs)-1}）"}

        para = doc.paragraphs[index]
        text_preview = para.text[:30]
        para._p.getparent().remove(para._p)

        _save(file_path)
        return {
            "success": True,
            "message": f"インデックス {index}「{text_preview}」を削除しました",
            "remaining_paragraphs": len(doc.paragraphs),
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def format_paragraph(file_path: str, index: int,
                     bold: Optional[bool] = None,
                     italic: Optional[bool] = None,
                     font_size: Optional[int] = None,
                     font_color: Optional[str] = None,
                     alignment: Optional[str] = None,
                     style: Optional[str] = None) -> dict:
    """段落全体の書式を変更する（alignment: left/center/right/justify）"""
    try:
        doc = _get_doc(file_path)
        if index < 0 or index >= len(doc.paragraphs):
            return {"success": False, "error": f"インデックス {index} は範囲外です"}

        para = doc.paragraphs[index]

        if style:
            para.style = doc.styles[style]

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if alignment and alignment.lower() in align_map:
            para.alignment = align_map[alignment.lower()]

        for run in para.runs:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if font_size is not None:
                run.font.size = Pt(font_size)
            if font_color is not None:
                hex_color = _validate_hex_color(font_color, "font_color")
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)

        _save(file_path)
        return {"success": True, "message": f"インデックス {index} の書式を更新しました"}
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def insert_image(file_path: str, image_path: str,
                 paragraph_index: Optional[int] = None,
                 width_cm: Optional[float] = None) -> dict:
    """画像をドキュメントに挿入する（paragraph_index省略時は末尾）"""
    try:
        if not os.path.exists(image_path):
            return {"success": False, "error": f"画像ファイルが見つかりません: {image_path}"}

        doc = _get_doc(file_path, create_if_missing=True)
        width = Cm(width_cm) if width_cm else None

        if paragraph_index is None:
            # 末尾に追加
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(image_path, width=width)
            inserted_index = len(doc.paragraphs) - 1
        else:
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return {"success": False, "error": f"インデックス {paragraph_index} は範囲外です"}
            # 指定段落の後に画像段落を挿入
            ref_para = doc.paragraphs[paragraph_index]
            new_para_elem = OxmlElement("w:p")
            ref_para._p.addnext(new_para_elem)
            inserted_index = paragraph_index + 1
            target_para = doc.paragraphs[inserted_index]
            run = target_para.add_run()
            run.add_picture(image_path, width=width)

        _save(file_path)
        return {
            "success": True,
            "message": f"画像「{os.path.basename(image_path)}」をインデックス {inserted_index} に挿入しました",
            "inserted_index": inserted_index,
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def add_table(file_path: str, data: list[list],
              paragraph_index: Optional[int] = None,
              has_header: bool = True) -> dict:
    """テーブルを追加する（data: 2次元配列、先頭行をヘッダーとして太字に）"""
    try:
        doc = _get_doc(file_path, create_if_missing=True)
        if not data:
            return {"success": False, "error": "data が空です"}

        rows = len(data)
        cols = max(len(row) for row in data)

        # テーブルを文書末尾に追加
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        for r_idx, row_data in enumerate(data):
            for c_idx, cell_val in enumerate(row_data):
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = str(cell_val) if cell_val is not None else ""
                if has_header and r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

        # 指定段落の後ろに移動
        if paragraph_index is not None and paragraph_index < len(doc.paragraphs):
            ref_para = doc.paragraphs[paragraph_index]
            tbl_elem = table._tbl
            tbl_elem.getparent().remove(tbl_elem)
            ref_para._p.addnext(tbl_elem)

        _save(file_path)
        return {
            "success": True,
            "message": f"{rows}行×{cols}列のテーブルを追加しました",
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def add_heading(file_path: str, text: str, level: int = 1) -> dict:
    """見出しを末尾に追加する（level: 1〜9）"""
    try:
        doc = _get_doc(file_path, create_if_missing=True)
        doc.add_heading(text, level=level)
        _save(file_path)
        idx = len(doc.paragraphs) - 1
        return {
            "success": True,
            "message": f"見出し(H{level})「{text}」をインデックス {idx} に追加しました",
            "index": idx,
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def save_word(file_path: str, save_as: Optional[str] = None) -> dict:
    """Wordファイルを保存する"""
    try:
        doc = _get_doc(file_path)
        target = os.path.abspath(save_as if save_as else file_path)
        doc.save(target)
        if save_as:
            _doc_cache[target] = doc
        return {"success": True, "message": f"'{target}' に保存しました"}
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def get_document_info(file_path: str) -> dict:
    """ドキュメントの構造情報（見出し・段落数・テーブル数）を返す"""
    try:
        doc = _get_doc(file_path)
        headings = [
            {"index": i, "level": int(p.style.name.split()[-1]) if p.style.name.startswith("Heading") else 0,
             "text": p.text}
            for i, p in enumerate(doc.paragraphs)
            if p.style.name.startswith("Heading")
        ]
        return {
            "success": True,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "headings": headings,
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def set_page_layout(file_path: str,
                    orientation: Optional[str] = None,
                    top_cm: Optional[float] = None,
                    bottom_cm: Optional[float] = None,
                    left_cm: Optional[float] = None,
                    right_cm: Optional[float] = None) -> dict:
    """ページの向き・余白を設定する（orientation: portrait / landscape）"""
    try:
        from docx.enum.section import WD_ORIENT
        doc = _get_doc(file_path, create_if_missing=True)
        section = doc.sections[0]

        if orientation is not None:
            if orientation.lower() in ("landscape", "横"):
                if section.orientation != WD_ORIENT.LANDSCAPE:
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width, section.page_height = section.page_height, section.page_width
            elif orientation.lower() in ("portrait", "縦"):
                if section.orientation != WD_ORIENT.PORTRAIT:
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width, section.page_height = section.page_height, section.page_width

        if top_cm    is not None: section.top_margin    = Cm(top_cm)
        if bottom_cm is not None: section.bottom_margin = Cm(bottom_cm)
        if left_cm   is not None: section.left_margin   = Cm(left_cm)
        if right_cm  is not None: section.right_margin  = Cm(right_cm)

        _save(file_path)
        return {"success": True, "message": "ページレイアウトを設定しました"}
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def add_page_break(file_path: str,
                   paragraph_index: Optional[int] = None) -> dict:
    """改ページを挿入する（paragraph_index 省略時は末尾）"""
    try:
        from docx.enum.text import WD_BREAK
        doc = _get_doc(file_path, create_if_missing=True)

        if paragraph_index is None:
            para = doc.add_paragraph()
            para.add_run().add_break(WD_BREAK.PAGE)
            inserted_index = len(doc.paragraphs) - 1
        else:
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return {"success": False, "error": f"インデックス {paragraph_index} は範囲外です"}
            ref_para = doc.paragraphs[paragraph_index]
            new_para_elem = OxmlElement("w:p")
            ref_para._p.addnext(new_para_elem)
            inserted_index = paragraph_index + 1
            doc.paragraphs[inserted_index].add_run().add_break(WD_BREAK.PAGE)

        _save(file_path)
        return {"success": True, "message": f"改ページをインデックス {inserted_index} に挿入しました",
                "inserted_index": inserted_index}
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def read_table(file_path: str, table_index: int) -> dict:
    """既存テーブルの内容を詳細に読み取る"""
    try:
        doc = _get_doc(file_path)
        if table_index < 0 or table_index >= len(doc.tables):
            return {"success": False,
                    "error": f"テーブルインデックス {table_index} は範囲外です（0〜{len(doc.tables)-1}）"}

        table = doc.tables[table_index]
        rows_data = [
            [{"row": r_idx, "col": c_idx, "text": cell.text}
             for c_idx, cell in enumerate(row.cells)]
            for r_idx, row in enumerate(table.rows)
        ]
        return {
            "success": True,
            "table_index": table_index,
            "row_count": len(table.rows),
            "col_count": len(table.columns),
            "rows": rows_data,
        }
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def format_table(file_path: str, table_index: int,
                 row: int, col: int,
                 bold: Optional[bool] = None,
                 italic: Optional[bool] = None,
                 font_size: Optional[int] = None,
                 font_color: Optional[str] = None,
                 bg_color: Optional[str] = None,
                 alignment: Optional[str] = None) -> dict:
    """テーブルの特定セルに書式を適用する"""
    try:
        doc = _get_doc(file_path)
        if table_index < 0 or table_index >= len(doc.tables):
            return {"success": False, "error": f"テーブルインデックス {table_index} は範囲外です"}

        table = doc.tables[table_index]
        if row < 0 or row >= len(table.rows):
            return {"success": False, "error": f"行インデックス {row} は範囲外です（0〜{len(table.rows)-1}）"}
        if col < 0 or col >= len(table.columns):
            return {"success": False, "error": f"列インデックス {col} は範囲外です（0〜{len(table.columns)-1}）"}

        cell = table.rows[row].cells[col]
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }

        for para in cell.paragraphs:
            if alignment and alignment.lower() in align_map:
                para.alignment = align_map[alignment.lower()]
            for run in para.runs:
                if bold      is not None: run.bold   = bold
                if italic    is not None: run.italic = italic
                if font_size is not None: run.font.size = Pt(font_size)
                if font_color is not None:
                    hc = _validate_hex_color(font_color, "font_color")
                    run.font.color.rgb = RGBColor(int(hc[0:2], 16), int(hc[2:4], 16), int(hc[4:6], 16))

        if bg_color is not None:
            hc = _validate_hex_color(bg_color, "bg_color")
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hc)
            tcPr.append(shd)

        _save(file_path)
        return {"success": True, "message": f"テーブル {table_index} の [{row}][{col}] に書式を適用しました"}
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


def add_header_footer(file_path: str,
                      header_text: Optional[str] = None,
                      footer_text: Optional[str] = None) -> dict:
    """ヘッダー・フッターを設定する"""
    try:
        doc = _get_doc(file_path, create_if_missing=True)
        section = doc.sections[0]

        if header_text is not None:
            header = section.header
            para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            para.text = header_text

        if footer_text is not None:
            footer = section.footer
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.text = footer_text

        _save(file_path)
        parts = []
        if header_text is not None: parts.append(f"ヘッダー「{header_text}」")
        if footer_text is not None: parts.append(f"フッター「{footer_text}」")
        return {"success": True, "message": f"{' / '.join(parts)} を設定しました"}
    except Exception as e:
        _reraise_if_fatal(e)
        return {"success": False, "error": str(e)}


# ── ツール定義（Claude API用）──────────────────────────────────────────────

TOOLS = [
    {
        "name": "open_word",
        "description": "Open a Word document. If the file does not exist, set create_if_missing=true to create it.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file (.docx)"},
                "create_if_missing": {"type": "boolean", "description": "Create the file if it does not exist"}
            },
            "required": ["file_path"]
        }
    },
    {
        "name": "read_document",
        "description": "Read the full content of a Word document with paragraph indices.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "include_tables": {"type": "boolean", "description": "Include table contents (default: true)"}
            },
            "required": ["file_path"]
        }
    },
    {
        "name": "read_paragraph",
        "description": "Read the text and formatting of a specific paragraph by index.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "index": {"type": "integer", "description": "Paragraph index (0-based)"}
            },
            "required": ["file_path", "index"]
        }
    },
    {
        "name": "append_paragraph",
        "description": "Append a new paragraph at the end of the document.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "text": {"type": "string", "description": "Text content to append"},
                "style": {"type": "string", "description": "Paragraph style, e.g. 'Normal' or 'Heading 1'"},
                "bold": {"type": "boolean", "description": "Set bold"},
                "italic": {"type": "boolean", "description": "Set italic"},
                "font_size": {"type": "integer", "description": "Font size in pt"}
            },
            "required": ["file_path", "text"]
        }
    },
    {
        "name": "insert_paragraph",
        "description": "Insert a paragraph before or after a specified paragraph index.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "index": {"type": "integer", "description": "Reference paragraph index"},
                "text": {"type": "string", "description": "Text to insert"},
                "position": {"type": "string", "enum": ["before", "after"], "description": "Insert before or after the reference paragraph (default: before)"},
                "style": {"type": "string", "description": "Paragraph style"},
                "bold": {"type": "boolean", "description": "Set bold"},
                "italic": {"type": "boolean", "description": "Set italic"},
                "font_size": {"type": "integer", "description": "Font size in pt"}
            },
            "required": ["file_path", "index", "text"]
        }
    },
    {
        "name": "replace_text",
        "description": "Find and replace text in the document. Useful for editing and proofreading.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "old_text": {"type": "string", "description": "Text to search for"},
                "new_text": {"type": "string", "description": "Replacement text"},
                "all_occurrences": {"type": "boolean", "description": "Replace all occurrences (default: true)"}
            },
            "required": ["file_path", "old_text", "new_text"]
        }
    },
    {
        "name": "delete_paragraph",
        "description": "Delete the paragraph at the specified index.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "index": {"type": "integer", "description": "Index of the paragraph to delete"}
            },
            "required": ["file_path", "index"]
        }
    },
    {
        "name": "format_paragraph",
        "description": "Apply formatting to a paragraph: bold, italic, font size, color, alignment, or style.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "index": {"type": "integer", "description": "Target paragraph index"},
                "bold": {"type": "boolean", "description": "Set bold"},
                "italic": {"type": "boolean", "description": "Set italic"},
                "font_size": {"type": "integer", "description": "Font size in pt"},
                "font_color": {"type": "string", "description": "Font color as hex string, e.g. FF0000 for red"},
                "alignment": {"type": "string", "enum": ["left", "center", "right", "justify"], "description": "Text alignment"},
                "style": {"type": "string", "description": "Paragraph style name, e.g. 'Heading 1' or 'Normal'"}
            },
            "required": ["file_path", "index"]
        }
    },
    {
        "name": "insert_image",
        "description": "Insert an image file into the document.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "image_path": {"type": "string", "description": "Path to the image file (PNG, JPG, etc.)"},
                "paragraph_index": {"type": "integer", "description": "Insert after this paragraph index (default: append at end)"},
                "width_cm": {"type": "number", "description": "Image width in centimeters (default: original size)"}
            },
            "required": ["file_path", "image_path"]
        }
    },
    {
        "name": "add_table",
        "description": "Add a table to the document from a 2D data array.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "data": {
                    "type": "array",
                    "items": {"type": "array"},
                    "description": "2D array of table data (list of rows)"
                },
                "paragraph_index": {"type": "integer", "description": "Insert after this paragraph index (default: append at end)"},
                "has_header": {"type": "boolean", "description": "Bold the first row as a header (default: true)"}
            },
            "required": ["file_path", "data"]
        }
    },
    {
        "name": "add_heading",
        "description": "Append a heading to the end of the document.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "text": {"type": "string", "description": "Heading text"},
                "level": {"type": "integer", "description": "Heading level 1-9 (default: 1)"}
            },
            "required": ["file_path", "text"]
        }
    },
    {
        "name": "save_word",
        "description": "Save the Word document. Optionally save to a new path.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "save_as": {"type": "string", "description": "New path to save a copy (omit to overwrite original)"}
            },
            "required": ["file_path"]
        }
    },
    {
        "name": "get_document_info",
        "description": "Get structural info about a document: heading list, paragraph count, and table count.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"}
            },
            "required": ["file_path"]
        }
    },
    {
        "name": "set_page_layout",
        "description": "Set page orientation (portrait/landscape) and margins.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "orientation": {"type": "string", "enum": ["portrait", "landscape"], "description": "Page orientation"},
                "top_cm":    {"type": "number", "description": "Top margin in cm"},
                "bottom_cm": {"type": "number", "description": "Bottom margin in cm"},
                "left_cm":   {"type": "number", "description": "Left margin in cm"},
                "right_cm":  {"type": "number", "description": "Right margin in cm"}
            },
            "required": ["file_path"]
        }
    },
    {
        "name": "add_page_break",
        "description": "Insert a page break after a paragraph index, or at the end if omitted.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "paragraph_index": {"type": "integer", "description": "Insert after this paragraph index (default: append at end)"}
            },
            "required": ["file_path"]
        }
    },
    {
        "name": "read_table",
        "description": "Read the contents of an existing table by index.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path":    {"type": "string",  "description": "Path to the Word file"},
                "table_index":  {"type": "integer", "description": "Table index (0-based)"}
            },
            "required": ["file_path", "table_index"]
        }
    },
    {
        "name": "format_table",
        "description": "Apply formatting to a specific cell in an existing table.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path":    {"type": "string",  "description": "Path to the Word file"},
                "table_index":  {"type": "integer", "description": "Table index (0-based)"},
                "row":          {"type": "integer", "description": "Row index (0-based)"},
                "col":          {"type": "integer", "description": "Column index (0-based)"},
                "bold":         {"type": "boolean"},
                "italic":       {"type": "boolean"},
                "font_size":    {"type": "integer", "description": "Font size in pt"},
                "font_color":   {"type": "string",  "description": "Font color as hex, e.g. FF0000"},
                "bg_color":     {"type": "string",  "description": "Background color as hex, e.g. FFFF00"},
                "alignment":    {"type": "string",  "enum": ["left", "center", "right", "justify"]}
            },
            "required": ["file_path", "table_index", "row", "col"]
        }
    },
    {
        "name": "add_header_footer",
        "description": "Set the header and/or footer text for the document.",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path":    {"type": "string", "description": "Path to the Word file"},
                "header_text":  {"type": "string", "description": "Header text (omit to leave unchanged)"},
                "footer_text":  {"type": "string", "description": "Footer text (omit to leave unchanged)"}
            },
            "required": ["file_path"]
        }
    },
    {
        "name": "append_rich_paragraph",
        "description": (
            "Append a paragraph built from a list of runs with individual bold/italic formatting. "
            "Use the 'runs' output from parse_markdown or parse_inline_formatting to preserve "
            "Markdown bold (**text**) and italic (*text*) as real Word formatting."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Path to the Word file"},
                "runs": {
                    "type": "array",
                    "description": "List of run objects: [{\"text\": str, \"bold\": bool, \"italic\": bool, \"color\": \"FF0000\"}]",
                    "items": {
                        "type": "object",
                        "properties": {
                            "text":   {"type": "string"},
                            "bold":   {"type": "boolean"},
                            "italic": {"type": "boolean"},
                            "color":     {"type": "string",  "description": "Font color as 6-digit hex, e.g. FF0000"},
                            "font_size": {"type": "integer", "description": "Font size in pt for this run only"},
                        },
                        "required": ["text"],
                    },
                },
                "style":     {"type": "string",  "description": "Paragraph style, e.g. 'Normal'"},
                "font_size": {"type": "integer", "description": "Font size in pt applied to all runs"},
            },
            "required": ["file_path", "runs"]
        }
    },
]

TOOL_FUNCTIONS = {
    "open_word": open_word,
    "read_document": read_document,
    "read_paragraph": read_paragraph,
    "append_paragraph": append_paragraph,
    "append_rich_paragraph": append_rich_paragraph,
    "insert_paragraph": insert_paragraph,
    "replace_text": replace_text,
    "delete_paragraph": delete_paragraph,
    "format_paragraph": format_paragraph,
    "insert_image": insert_image,
    "add_table": add_table,
    "add_heading": add_heading,
    "save_word": save_word,
    "get_document_info": get_document_info,
    "set_page_layout": set_page_layout,
    "add_page_break": add_page_break,
    "read_table": read_table,
    "format_table": format_table,
    "add_header_footer": add_header_footer,
}


def execute_tool(tool_name: str, tool_input: dict) -> str:
    func = TOOL_FUNCTIONS.get(tool_name)
    if not func:
        return json.dumps({"success": False, "error": f"不明なツール: {tool_name}"}, ensure_ascii=False)
    result = func(**tool_input)
    return json.dumps(result, ensure_ascii=False, indent=2)
